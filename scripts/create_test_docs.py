"""Generate sample DOCX files for migration smoke testing.

Creates two NDA-style documents in the project root:
  - test_contract.docx        (the original)
  - test_contract_v2.docx     (a slightly modified version, used by /compare-documents)
"""

from docx import Document


_DOC1_PARAGRAPHS = [
    'This Non-Disclosure Agreement (the "Agreement") is entered into as of January 1, 2025, '
    'between Acme Corporation, a Delaware corporation ("Disclosing Party"), and '
    'John Smith, an individual ("Receiving Party").',
    '1. Definition of Confidential Information. '
    '"Confidential Information" shall mean any non-public information disclosed by either party, '
    'including trade secrets, business plans, customer lists, technical data, and product designs.',
    '2. Obligations of Receiving Party. '
    'The Receiving Party agrees to maintain the confidentiality of all Confidential Information '
    'and shall not disclose it to any third party without prior written consent of the Disclosing Party.',
    '3. Term. '
    'This Agreement shall remain in effect for a period of three (3) years from the Effective Date.',
    '4. Governing Law. '
    'This Agreement shall be governed by and construed in accordance with the laws of the State of Delaware.',
    '5. Termination. '
    'Either party may terminate this Agreement at any time with thirty (30) days written notice to the other party.',
    '6. Remedies. '
    'The Receiving Party acknowledges that money damages may not be a sufficient remedy for breach of this Agreement '
    'and that the Disclosing Party shall be entitled to seek injunctive relief.',
]

_DOC2_PARAGRAPHS = [
    'This Non-Disclosure Agreement (the "Agreement") is entered into as of January 1, 2025, '
    'between Acme Corporation, a Delaware corporation ("Disclosing Party"), and '
    'John Smith, an individual ("Receiving Party").',
    '1. Definition of Confidential Information. '
    '"Confidential Information" shall mean any non-public information disclosed by either party, '
    'including trade secrets, business plans, customer lists, technical data, and product designs.',
    '2. Obligations of Receiving Party. '
    'The Receiving Party agrees to maintain the confidentiality of all Confidential Information '
    'and shall not disclose it to any third party without prior written consent of the Disclosing Party. '
    'The Receiving Party shall return all Confidential Information upon termination of this Agreement.',
    '3. Term. '
    'This Agreement shall remain in effect for a period of five (5) years from the Effective Date.',
    '4. Governing Law. '
    'This Agreement shall be governed by and construed in accordance with the laws of the State of California.',
    '5. Termination. '
    'Either party may terminate this Agreement at any time with sixty (60) days written notice to the other party.',
    '6. Remedies. '
    'The Receiving Party acknowledges that money damages may not be a sufficient remedy for breach of this Agreement '
    'and that the Disclosing Party shall be entitled to seek injunctive relief and recover attorneys fees.',
]


def _build(paragraphs: list[str], path: str) -> None:
    doc = Document()
    doc.add_heading("NON-DISCLOSURE AGREEMENT", level=0)
    for para in paragraphs:
        doc.add_paragraph(para)
    doc.save(path)
    print(f"Created {path}")


def main() -> None:
    _build(_DOC1_PARAGRAPHS, "test_contract.docx")
    _build(_DOC2_PARAGRAPHS, "test_contract_v2.docx")


if __name__ == "__main__":
    main()
