import io
import json
from typing import Any

from docx import Document
from fastapi import APIRouter, Depends, Header, HTTPException, UploadFile
from fastapi.responses import StreamingResponse

# from src.api.session_utils import get_session_id
from src.config.logging import get_logger
from src.core.auth import generate_access_token, get_session_id, verify_token
from src.schemas.contract_analyzer import ContractAnalyzerResponse
from src.schemas.doc_chat import DocChatResponse
from src.schemas.general_review import GeneralReviewRequest, GeneralReviewResponse
from src.schemas.playbook_review import (
    PlayBookReviewFinalResponse,
    RuleCheckRequest,
)
from src.tools.comparision import compare_documents_stream_service
from src.tools.comparision import run as compare_documents_service
from src.tools.doc_chat import query_document as query_document_service
from src.tools.general_review import clause_review, full_document_review
from src.tools.key_information import (
    get_key_information_document as contract_analyzer_service,
)
from src.tools.key_information import (
    stream_key_information as contract_analyzer_stream_service,
)
from src.tools.playbook_review import review_document as playbook_review_service

router = APIRouter(tags=["agents"])
logger = get_logger("agents")


@router.get("/generate-access-token")
def get_access_token() -> Any:
    """Generate an access token for the application."""

    access_token = generate_access_token()

    return access_token


@router.post("/compare-documents")
async def compare_documents_endpoint(file_a: UploadFile, file_b: UploadFile, session_id: str = Depends(get_session_id)) -> Any:
    """Compare two documents and return their differences."""

    document_a = Document(io.BytesIO(await file_a.read()))
    document_b = Document(io.BytesIO(await file_b.read()))

    comparison_result = await compare_documents_service(session_id=session_id, document_a=document_a, document_b=document_b)
    return comparison_result


@router.post("/compare-documents/stream", response_class=StreamingResponse)
async def compare_documents_stream_endpoint(file_a: UploadFile, file_b: UploadFile, session_id: str = Depends(get_session_id)) -> StreamingResponse:
    """Compare two documents and stream each change as its analysis lands (SSE).

    Emits one ``data: {json}`` frame per event (``status`` -> ``change`` per clause,
    fastest first -> final ``summary``), terminated by ``data: [DONE]``. The frontend
    renders each ``change`` as a card the moment it arrives instead of waiting for the
    whole comparison.
    """

    document_a = Document(io.BytesIO(await file_a.read()))
    document_b = Document(io.BytesIO(await file_b.read()))

    headers = {"Access-Control-Allow-Origin": "*", "Cache-Control": "no-cache", "Connection": "keep-alive", "X-Accel-Buffering": "no"}
    return StreamingResponse(
        compare_documents_stream_service(session_id=session_id, document_a=document_a, document_b=document_b),
        media_type="text/event-stream",
        headers=headers,
    )


@router.post("/playbook-review", response_model=PlayBookReviewFinalResponse)
async def playbook_review_endpoint(request: RuleCheckRequest, session_id: str = Depends(get_session_id)) -> PlayBookReviewFinalResponse:
    """Run playbook validation checks."""

    review_result = await playbook_review_service(session_id=session_id, request=request)
    return review_result


@router.post("/general-review", response_model=GeneralReviewResponse)
async def review_contract(request: GeneralReviewRequest, session_id: str = Depends(get_session_id)) -> GeneralReviewResponse:
    """Run the general review agent against an ingested document."""
    try:
        if request.selected_clause and request.selected_clause.strip():
            return await clause_review(
                session_id=session_id,
                clause_text=request.selected_clause,
                user_prompt=request.prompt,
                clause_title=(request.clause_title or "Selected Clause").strip() or "Selected Clause",
            )

        return await full_document_review(
            session_id=session_id,
            user_prompt=request.prompt,
        )

    except ValueError as err:
        raise HTTPException(status_code=400, detail=str(err))
    except Exception as err:
        raise HTTPException(status_code=500, detail=f"General review error: {str(err)}")


@router.post("/contract-analyzer")
async def contract_analyzer_endpoint(file: UploadFile, session_id: str = Depends(get_session_id)) -> ContractAnalyzerResponse:
    """Analyze a contract document and extract key information."""

    document = Document(io.BytesIO(await file.read()))
    document_data = "\n".join([para.text for para in document.paragraphs if para.text.strip() != ""])

    analysis_result: ContractAnalyzerResponse = await contract_analyzer_service(content=document_data, session_id=session_id)  # type: ignore
    return analysis_result


async def _sse_event_stream(event_source: Any) -> Any:
    """Format ``(event_name, data)`` tuples from the analyzer generator as SSE wire frames.

    Always emits a terminal ``error`` frame if the source raises, so the client never
    hangs on a dropped stream.
    """

    try:
        async for event_name, data in event_source:
            yield f"event: {event_name}\ndata: {json.dumps(data)}\n\n"
    except Exception as exc:
        logger.error(f"Contract analyzer stream failed: {exc}")
        yield f"event: error\ndata: {json.dumps({'message': 'Streaming failed. Please retry.'})}\n\n"


async def _with_document_event(document_text: str, source: Any) -> Any:
    """Demo-only: prepend a ``document`` event carrying the parsed contract text so the
    demo page can render the source document beside the streaming analysis. The production
    Word plugin already has the document open in Word, so it omits ``include_document``."""

    yield ("document", {"text": document_text})
    async for event in source:
        yield event


@router.post("/contract-analyzer/stream")
async def contract_analyzer_stream_endpoint(
    file: UploadFile,
    session_id: str = Depends(get_session_id),
    include_document: bool = False,
) -> Any:
    """Stream the contract analysis section-by-section (SSE) as each parallel section resolves.

    Emits ``start`` (the section names) -> one ``section`` event per section as it lands
    (fastest first, each a fully schema-validated payload) -> ``done``. The merged result is
    cached on the session, so the existing non-stream POST /contract-analyzer keeps working
    and returns instantly on a warm session.
    """

    document = Document(io.BytesIO(await file.read()))
    document_data = "\n".join([para.text for para in document.paragraphs if para.text.strip() != ""])

    source = contract_analyzer_stream_service(content=document_data, session_id=session_id)
    if include_document:
        source = _with_document_event(document_data, source)

    return StreamingResponse(
        _sse_event_stream(source),
        media_type="text/event-stream",
        headers={"X-Accel-Buffering": "no", "Cache-Control": "no-cache", "Connection": "keep-alive"},
    )


@router.post("/query-document", response_model=DocChatResponse)
async def query_document_endpoint(query: str, session_id: str = Depends(get_session_id)) -> DocChatResponse:
    """Query the document chunks based on the given query and session ID."""

    llm_result = await query_document_service(query=query, session_id=session_id)
    return llm_result
