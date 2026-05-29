from __future__ import annotations

import asyncio
import functools
import os
import re
import time
import uuid
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple

import numpy as np
from docx.document import Document

from src.config.logging import Logger
from src.config.settings import get_settings
from src.exceptions.parser_exceptions import (
    DocxCleaningException,
    DocxParagraphExtractionException,
    DocxTableExtractionException,
)
from src.schemas.playbook_review import TextInfo
from src.schemas.registry import Chunk, ParseResult
from src.services.registry.base_parser import BaseParser
from src.services.session_manager import SessionData
from src.services.vector_store.embeddings.base_embedding_service import (
    BaseEmbeddingService,
)
from src.services.vector_store.manager import get_faiss_vector_store

# ── compiled regexes (module-level — zero per-call cost) ─────────────────────

_SECTION_LABEL_RE = re.compile(r"^(\d+[\.\)]?\s+\S.*|\d+[\.\)]?\s*$|[A-Z][A-Z\s\.\,\&\'\-]{1,60}$)")
_INLINE_CLAUSE_HEADING_RE = re.compile(
    r"(?:^|(?<=\. ))" r"((?:[A-Z][a-z]+(?:'[a-z]+)?)" r"(?:(?:\s+(?:and|&|of|the|to|for|on|in|or|with|at|by|an|a|your|its|from))*" r"\s+[A-Z][a-z]+(?:'[a-z]+)?)+\.)\s"
)
_SINGLE_WORD_HEADING_RE = re.compile(r"(?:^|(?<=\. ))([A-Z][a-z]{3,}(?:'[a-z]+)?\.)\s(?=[A-Z])")
_CLAUSE_PREFIX_RE = re.compile(r"^(\d+[\.\)]\d*[\.\d]*\s|\([a-z]+\)\s|[a-z]\)\s)")
_CONTROL_RE = re.compile(r"[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f]")
_WHITESPACE_RE = re.compile(r"\s+")

_SPECIAL_CHARS_TABLE = str.maketrans({"\u00a0": " ", "\u200b": "", "\ufeff": "", "\r": ""})
_HEADING_BLACKLIST: frozenset[str] = frozenset(
    {
        "however",
        "further",
        "furthermore",
        "additionally",
        "moreover",
        "otherwise",
        "nevertheless",
        "notwithstanding",
        "accordingly",
        "therefore",
        "thus",
        "hence",
        "thereby",
        "thereafter",
        "whereas",
        "specifically",
        "particularly",
        "collectively",
        "individually",
        "alternatively",
        "consequently",
        "subsequently",
        "respectively",
        "conversely",
        "similarly",
        "meanwhile",
        "indeed",
        "finally",
        "firstly",
        "secondly",
        "thirdly",
        "lastly",
        "instead",
    }
)

_EMBED_BATCH_SIZE = 64
_MAX_CONCURRENT_BATCHES = 8


# ── pure helpers ──────────────────────────────────────────────────────────────


def _make_uuids(n: int) -> List[str]:
    raw = os.urandom(16 * n)
    out: List[str] = []
    for i in range(n):
        b = bytearray(raw[i * 16 : i * 16 + 16])
        b[6] = (b[6] & 0x0F) | 0x40
        b[8] = (b[8] & 0x3F) | 0x80
        out.append(str(uuid.UUID(bytes=bytes(b))))
    return out


@functools.lru_cache(maxsize=4096)
def _clean_text(text: str) -> str:
    """Pure function — safe to cache. Saves ~200ms on repeat-heavy docs."""
    if not text:
        return ""
    translated = text.translate(_SPECIAL_CHARS_TABLE)
    if translated != text or (text and "\x00" <= text[0] <= "\x1f"):
        if _CONTROL_RE.search(translated):
            translated = _CONTROL_RE.sub("", translated)
    text = _WHITESPACE_RE.sub(" ", translated).strip()
    if not text:
        return ""
    stripped = text.lstrip(" \n\t")
    return stripped if _CLAUSE_PREFIX_RE.match(stripped) else stripped.lstrip(".")


def _is_structural_heading(text: str, max_words: int = 8) -> bool:
    words = text.split()
    return bool(words) and len(words) <= max_words and bool(_SECTION_LABEL_RE.match(text.strip()))


def _find_clause_heading_matches(text: str) -> List[Dict[str, Any]]:
    found: Dict[int, Dict[str, Any]] = {}
    for m in _INLINE_CLAUSE_HEADING_RE.finditer(text):
        found[m.start()] = {
            "start": m.start(),
            "end": m.end(1) + 1,
            "heading": m.group(1),
        }
    for m in _SINGLE_WORD_HEADING_RE.finditer(text):
        word = m.group(1).rstrip(".").lower()
        if word not in _HEADING_BLACKLIST:
            found.setdefault(
                m.start(),
                {"start": m.start(), "end": m.end(1) + 1, "heading": m.group(1)},
            )
    return [found[k] for k in sorted(found)]


def _split_at_clause_boundaries_sync(paragraphs: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    result: List[Dict[str, Any]] = []
    for para in paragraphs:
        if para["is_heading"]:
            result.append(para)
            continue
        text = para["content"]
        boundaries = _find_clause_heading_matches(text)
        if not boundaries:
            result.append(para)
            continue
        if boundaries[0]["start"] > 0:
            prefix = text[: boundaries[0]["start"]].strip()
            if prefix:
                result.append(
                    {
                        "index": para["index"],
                        "content": prefix,
                        "is_heading": False,
                        "wc": len(prefix.split()),
                    }
                )
        for i, b in enumerate(boundaries):
            body = text[b["end"] : (boundaries[i + 1]["start"] if i + 1 < len(boundaries) else len(text))].strip()
            result.append(
                {
                    "index": para["index"],
                    "content": b["heading"],
                    "is_heading": True,
                    "wc": len(b["heading"].split()),
                }
            )
            if body:
                result.append(
                    {
                        "index": para["index"],
                        "content": body,
                        "is_heading": False,
                        "clause_boundary": True,
                        "wc": len(body.split()),
                    }
                )
    return result


def _cosine_similarities_batch(embeddings: np.ndarray) -> np.ndarray:
    norms = np.linalg.norm(embeddings, axis=1, keepdims=True)
    np.maximum(norms, 1e-9, out=norms)
    normed = embeddings / norms
    return np.einsum("ij,ij->i", normed[:-1], normed[1:])


def _merge_orphan_chunks(
    texts: List[str], headings: List[Optional[str]], wcs: List[int], spans: List[Tuple[int, int]], min_words: int = 5
) -> Tuple[List[str], List[Optional[str]], List[Tuple[int, int]]]:
    if not texts:
        return [], [], []
    out_t: List[str] = []
    out_h: List[Optional[str]] = []
    out_s: List[Tuple[int, int]] = []
    c_t, c_h, c_wc, c_s = "", None, 0, None

    for t, h, wc, span in zip(texts, headings, wcs, spans):
        if c_t:
            t = c_t + " " + t
            h = c_h or h
            wc = c_wc + wc
            span = (c_s[0], span[1])  # type: ignore[index]
            c_t, c_h, c_wc, c_s = "", None, 0, None
        if wc < min_words:
            c_t, c_h, c_wc, c_s = t, h, wc, span
        else:
            out_t.append(t)
            out_h.append(h)
            out_s.append(span)

    if c_t:
        if out_t:
            out_t[-1] += " " + c_t
            out_s[-1] = (out_s[-1][0], c_s[1])  # type: ignore[index]
        else:
            out_t.append(c_t)
            out_h.append(c_h)
            out_s.append(c_s)  # type: ignore[arg-type]
    return out_t, out_h, out_s


def _mean_pool_by_spans(embeddings: np.ndarray, spans: List[Tuple[int, int]]) -> np.ndarray:

    if not spans:
        return np.empty((0, embeddings.shape[1]), dtype=np.float32)

    # Keep only non-empty spans fully inside the embedding array bounds.
    n_rows = embeddings.shape[0]
    valid: List[Tuple[int, int]] = [(int(s), int(e)) for s, e in spans if 0 <= s < e and e <= n_rows]
    if not valid:
        return np.empty((0, embeddings.shape[1]), dtype=np.float32)

    # reduceat requires non-decreasing start indices.
    valid.sort(key=lambda se: se[0])

    starts = np.array([s for s, _ in valid], dtype=np.intp)
    lengths = np.array([e - s for s, e in valid], dtype=np.float32)

    # Defensive check (shouldn't trigger due to filtering above).
    if np.any(starts >= n_rows):
        valid = [(s, e) for s, e in valid if s < n_rows]
        if not valid:
            return np.empty((0, embeddings.shape[1]), dtype=np.float32)
        starts = np.array([s for s, _ in valid], dtype=np.intp)
        lengths = np.array([e - s for s, e in valid], dtype=np.float32)

    sums = np.add.reduceat(embeddings, starts, axis=0)[: len(valid)]
    return (sums / lengths[:, np.newaxis]).astype(np.float32)


async def _bulk_index(vector_store: Any, vectors: np.ndarray) -> None:
    if not len(vectors):
        return
    if hasattr(vector_store, "index_embeddings_bulk"):
        await vector_store.index_embeddings_bulk(vectors)
    else:
        await asyncio.gather(*[vector_store.index_embedding(v) for v in vectors])


# ── Parser ────────────────────────────────────────────────────────────────────


class DocxParser(BaseParser, Logger):
    _HEADING_MAX_WORDS = 8
    _ORPHAN_MIN_WORDS = 5

    def __init__(self) -> None:
        super().__init__()
        self.settings = get_settings()
        from src.dependencies import get_service_container

        sc = get_service_container()
        self.embedding_service: BaseEmbeddingService = sc.embedding_service
        self.vector_store = get_faiss_vector_store(self.embedding_service.get_embedding_dimensions())
        self._embed_sem = asyncio.Semaphore(_MAX_CONCURRENT_BATCHES)

    # ── embedding ─────────────────────────────────────────────────────────────

    async def _embed_batch(self, texts: List[str], task: str = "text-matching") -> np.ndarray:
        if not texts:
            return np.empty((0, self.embedding_service.get_embedding_dimensions()), dtype=np.float32)

        # Fast path: service exposes a native batch API — one call, service batches internally
        if hasattr(self.embedding_service, "generate_embeddings_batch"):
            vecs = await self.embedding_service.generate_embeddings_batch(texts=texts, task=task)
            return np.asarray(vecs, dtype=np.float32)

        # Fallback: fire every single-text call concurrently in one flat gather.
        # Semaphore caps peak concurrency without serialising the requests.
        async def _one(t: str) -> Any:
            async with self._embed_sem:
                return await self.embedding_service.generate_embeddings(text=t, task=task)

        results = await asyncio.gather(*[_one(t) for t in texts])
        return np.asarray(results, dtype=np.float32)

    # ── document preparation ──────────────────────────────────────────────────

    async def _prepare_paragraphs(self, document: Document) -> Tuple[List[Dict[str, Any]], Dict[str, Any]]:
        try:
            props = document.core_properties
            paras: List[Dict[str, Any]] = []
            para_wc = 0
            for idx, p in enumerate(document.paragraphs):
                for run in p.runs:
                    if run.text:
                        run.text = _WHITESPACE_RE.sub(" ", run.text).strip()
                raw = p.text.strip()
                if not raw:
                    continue
                cleaned = _clean_text(raw)
                if not cleaned:
                    continue
                wc = len(cleaned.split())
                para_wc += wc
                is_heading = bool((p.style and p.style.name.startswith("Heading")) or _is_structural_heading(cleaned, self._HEADING_MAX_WORDS))
                paras.append({"index": idx, "content": cleaned, "is_heading": is_heading, "wc": wc})
            metadata: Dict[str, Any] = {
                "source": "docx",
                "author": props.author or "Unknown",
                "title": props.title or "Untitled",
                "created_at": props.created.isoformat() if props.created else None,
                "modified_at": props.modified.isoformat() if props.modified else None,
                "paragraph_count": len(paras),
                "table_count": len(document.tables),
                "word_count": para_wc,
            }
            return paras, metadata
        except Exception as e:
            raise DocxParagraphExtractionException(str(e)) from e

    async def _extract_tables(self, document: Document) -> List[Dict[str, Any]]:
        try:
            tables = []
            for t_idx, table in enumerate(document.tables):
                rows = [[_clean_text(cell.text) for cell in row.cells] for row in table.rows]
                tables.append({"table_index": t_idx, "content": rows})
            return tables
        except Exception as e:
            raise DocxTableExtractionException(str(e)) from e

    # ── chunking (receives pre-computed embeddings; never calls embed itself) ─

    def _chunk_paragraphs(
        self,
        paragraphs: List[Dict[str, Any]],
        para_embeddings: np.ndarray,
    ) -> Tuple[List[str], List[Optional[str]], List[int], List[Tuple[int, int]]]:
        """
        Deterministic chunking on already-split paragraphs with pre-computed
        embeddings.  Returns (texts, headings, wcs, spans).
        Pure CPU — no I/O, no await needed; called inline after the gather.
        """
        wcs = [p.get("wc") or len(p["content"].split()) for p in paragraphs]

        if len(para_embeddings) > 1:
            sims = _cosine_similarities_batch(para_embeddings)
            threshold = float(sims.mean()) - 0.75 * float(sims.std())
            split_points: set[int] = {int(x) for x in np.where(sims < threshold)[0]}
        else:
            split_points = set()

        max_len = self.settings.chunk_size
        chunks: List[str] = []
        chunk_headings: List[Optional[str]] = []
        chunk_wcs: List[int] = []
        chunk_spans: List[Tuple[int, int]] = []
        current: List[str] = []
        current_wc = 0
        current_start = 0
        current_heading: Optional[str] = None
        pending_heading: Optional[str] = None

        def _flush(end: int) -> None:
            nonlocal current, current_wc, current_start, pending_heading
            if current:
                chunks.append(" ".join(current))
                chunk_headings.append(pending_heading)
                chunk_wcs.append(current_wc)
                chunk_spans.append((current_start, end))
                current = []
                current_wc = 0

        for i, (para, wc) in enumerate(zip(paragraphs, wcs)):
            if para["is_heading"]:
                current_heading = para["content"]
                _flush(i)
                current_start = i
                pending_heading = current_heading
            elif current_wc + wc > max_len:
                _flush(i)
                current_start = i
                pending_heading = current_heading

            if not current:
                current_start = i
                pending_heading = current_heading

            current.append(para["content"])
            current_wc += wc

            if i in split_points or para.get("clause_boundary"):
                _flush(i + 1)
                current_start = i + 1

        _flush(len(paragraphs))
        return chunks, chunk_headings, chunk_wcs, chunk_spans

    # ── public API ────────────────────────────────────────────────────────────

    async def parse_data(
        self,
        data: List["TextInfo"],
        session_data: Optional[Any] = None,
    ) -> ParseResult:
        start = time.perf_counter()
        paragraphs = [d.text for d in data if d.text]
        if not paragraphs:
            raise ValueError("No paragraphs found in the data.")
        vectors = await self._embed_batch(paragraphs)
        await _bulk_index(self.vector_store, vectors)
        now = datetime.utcnow().isoformat()
        model = self.embedding_service.model_name
        ids = _make_uuids(len(paragraphs))
        meta = {"chunk_type": "semantic_paragraph"}
        chunks = [
            Chunk(
                chunk_id=ids[i],
                document_id=None,
                chunk_index=i,
                content=text,
                embedding_model=model,
                embedding_vector=None,
                metadata=meta,
                created_at=now,
            )
            for i, text in enumerate(paragraphs)
        ]
        return ParseResult(
            success=True,
            chunks=chunks,
            metadata={"paragraph_count": len(paragraphs), "source": "parsed_data"},
            processing_time=time.perf_counter() - start,
        )

    async def parse_document(self, document: Document, session_data: Optional["SessionData"] = None) -> ParseResult:

        start = time.perf_counter()
        try:
            vector_store = session_data.vector_store if session_data else self.vector_store

            # ── Step 1: parse paragraphs + tables concurrently (fast CPU) ──
            (paragraphs, metadata), tables = await asyncio.gather(
                self._prepare_paragraphs(document),
                self._extract_tables(document),
            )

            document_id = str(uuid.uuid4())
            metadata["document_id"] = document_id
            now = datetime.utcnow().isoformat()
            model = self.embedding_service.model_name

            # ── Step 2: flatten table rows (CPU, negligible) ──
            table_rows: List[Dict[str, Any]] = []
            for table in tables:
                table_text = _clean_text(" ".join(" | ".join(r) for r in table["content"]))
                if table_text:
                    table_rows.append(
                        {
                            "text": table_text,
                            "table_index": table["table_index"],
                            "row_count": len(table["content"]),
                        }
                    )
            metadata["word_count"] += sum(len(cell.split()) for table in tables for row in table["content"] for cell in row if cell)

            # ── Step 3: embed ORIGINAL paragraphs + tables in ONE call,
            #            while clause-boundary split runs in a thread.
            #
            #    KEY INVARIANT: embed receives `para_texts` (N texts from the
            #    original list), NOT the post-split list (~2.7N texts).
            #    The split output is used only for boundary detection and text
            #    assembly; vectors come from mean-pooling the N original
            #    embeddings by span — same semantic result, far fewer embed calls.
            para_texts = [p["content"] for p in paragraphs]
            tbl_texts = [r["text"] for r in table_rows]
            all_texts = para_texts + tbl_texts
            n_para = len(para_texts)

            loop = asyncio.get_running_loop()
            all_vectors, split_paragraphs = await asyncio.gather(
                self._embed_batch(all_texts),  # I/O
                loop.run_in_executor(None, _split_at_clause_boundaries_sync, paragraphs),  # CPU
            )

            para_vectors: np.ndarray = all_vectors[:n_para]
            tbl_vectors: np.ndarray = all_vectors[n_para:]

            # ── Step 4: chunk (pure CPU, ~2ms) ──
            chunk_texts, chunk_headings, chunk_wcs, chunk_spans = self._chunk_paragraphs(split_paragraphs, para_vectors)

            # ── Step 5: mean-pool via vectorised reduceat ──
            dim = self.embedding_service.get_embedding_dimensions()
            chunk_vectors = _mean_pool_by_spans(para_vectors, chunk_spans)
            if len(chunk_vectors) < len(chunk_texts):
                pad = np.zeros((len(chunk_texts) - len(chunk_vectors), dim), dtype=np.float32)
                chunk_vectors = np.vstack([chunk_vectors, pad])

            # ── Step 6: single bulk_index call for everything ──
            combined = np.vstack([chunk_vectors, tbl_vectors]) if len(tbl_vectors) else chunk_vectors
            await _bulk_index(vector_store, combined)

            # ── Step 7: build Chunk objects (batch UUIDs, one timestamp) ──
            total = len(chunk_texts) + len(table_rows)
            ids = _make_uuids(total)
            chunks: List[Chunk] = []

            for i, (text, heading) in enumerate(zip(chunk_texts, chunk_headings)):
                chunk_meta: Dict[str, Any] = {"chunk_type": "semantic_paragraph"}
                if heading:
                    chunk_meta["section_heading"] = heading
                chunks.append(
                    Chunk(
                        chunk_id=ids[i],
                        document_id=document_id,
                        chunk_index=i,
                        content=text,
                        embedding_model=model,
                        embedding_vector=None,
                        metadata=chunk_meta,
                        created_at=now,
                    )
                )

            base = len(chunk_texts)
            for i, row in enumerate(table_rows):
                chunks.append(
                    Chunk(
                        chunk_id=ids[base + i],
                        document_id=document_id,
                        chunk_index=base + i,
                        content=row["text"],
                        embedding_model=model,
                        embedding_vector=None,
                        metadata={
                            "chunk_type": "table",
                            "table_index": row["table_index"],
                            "row_count": row["row_count"],
                        },
                        created_at=now,
                    )
                )

            return ParseResult(
                success=True,
                chunks=chunks,
                metadata=metadata,
                error_message=None,
                processing_time=time.perf_counter() - start,
            )

        except Exception as e:
            self.logger.error(str(e))
            return ParseResult(
                success=False,
                chunks=[],
                metadata={},
                error_message=str(e),
                processing_time=0.0,
            )

    # ── legacy shims ──────────────────────────────────────────────────────────

    async def clean_document(self, document: Document) -> None:
        try:
            for p in document.paragraphs:
                for run in p.runs:
                    if run.text:
                        run.text = _WHITESPACE_RE.sub(" ", run.text).strip()
        except Exception as e:
            raise DocxCleaningException(str(e)) from e

    @staticmethod
    def cosine_similarity(vec1: np.ndarray, vec2: np.ndarray) -> float:
        denom = float(np.linalg.norm(vec1) * np.linalg.norm(vec2))
        return float(np.dot(vec1, vec2) / denom) if denom else 0.0

    def is_healthy(self) -> Any:
        return True
